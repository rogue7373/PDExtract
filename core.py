import pandas as pd 
from docx.api import Document
import pathlib as Path
import os

username = os.getlogin() #Currently, this is not being used. 

document = Document(input('Enter file name...including the file extension here... '))

# Inserted to help with development. 
#document = Document('C:\\Users||jeftaylor\\Downloads\\INCIDENT_+UK+Buy+and+Appeals+Teammates+at_+CNX+Baguio+Unable+to+Connect+to+PHones.docx')


tables = []

wanted_tables = [0, 1, 2, 4, 5, 6, 7, 11, 12, 13, 15, 17]

for index,table in enumerate(document.tables):
    if index in wanted_tables:
        df = [['' for i in range(len(table.column))] for j in range(len(table.rows))]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                df[i][j] = cell.txt
        tables.extend(df)


        #print("generated table") # Inserted to help with development; so you know that tables are being captured properly.
    else: 
        continue


        #print("skipped") # Inserted to help with development; so you know that tables are being skipped properly.

# This can be uncommented to help with development or testing. 
#print(tables)

pd.DataFrame(tables).T.to_excel('tables.xlsx')

